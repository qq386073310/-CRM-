from PyQt5.QtWidgets import (QWidget, QVBoxLayout, QHBoxLayout, QLabel, 
                             QPushButton, QListWidget, QListWidgetItem,
                             QLineEdit, QTextEdit, QFormLayout, QDialog, QMessageBox, 
                             QComboBox, QDateEdit, QDoubleSpinBox, QFrame, 
                             QAbstractItemView, QCheckBox, QGridLayout, QMenu,
                             QTabWidget, QTableWidget, QTableWidgetItem, QHeaderView, QFileDialog,
                             QInputDialog, QGroupBox, QScrollArea, QSizePolicy)
from PyQt5.QtCore import Qt, QDate, QSize, pyqtSignal
from PyQt5.QtGui import QIcon, QDesktopServices, QIntValidator
from PyQt5.QtCore import QUrl
try:
    from docxtpl import DocxTemplate
except ImportError:
    DocxTemplate = None
import os
import shutil
from datetime import datetime
from core.logger import logger
from core.utils import get_app_path
from core.async_utils import Worker, QThreadPool
from modules.common_widgets import SingleSelectionWidget, ModernDateEdit
from modules.base_card import BaseCardWidget
from core.constants import CONTRACT_STATUS_MAP


class ContractCardWidget(BaseCardWidget):
    """合同卡片控件"""
    def __init__(self, data, parent=None):
        super().__init__(data, parent)
        self.setProperty("class", "card contract-card")
        # 设置固定大小，保持卡片一致性
        self.setFixedSize(360, 250)
        

    def _init_ui(self):
        # 设置提示信息（备注）
        remarks = self.data.get('remarks', '')
        if remarks:
            self.setToolTip(f"备注: {remarks}")
            
        self.setProperty("class", "card")
        
        # Use existing main_layout from BaseCardWidget
        # main_layout = QVBoxLayout(self) 
        main_layout = self.main_layout
        main_layout.setContentsMargins(15, 15, 15, 15)
        main_layout.setSpacing(8)
        
        # 1. 顶部：合同名称 + 编号
        top_layout = QHBoxLayout()
        
        top_layout.addWidget(self.create_checkbox())
        
        title_layout = QVBoxLayout()
        title_layout.setSpacing(2)
        
        # 合同名称
        title_label = QLabel(self.data.get('title', '无标题'))
        title_label.setStyleSheet("font-size: 14px; font-weight: bold;")
        title_label.setWordWrap(True)
        title_layout.addWidget(title_label)
        
        # 合同编号
        number_label = QLabel(self.data.get('contract_number', ''))
        number_label.setProperty("class", "info-text")
        number_label.setStyleSheet("font-size: 11px; color: #909399;")
        title_layout.addWidget(number_label)
        
        top_layout.addLayout(title_layout, 1)
        
        # 状态标签
        status = self.data.get('status', 'draft')
        status_text, status_color = CONTRACT_STATUS_MAP.get(status, ('未知', '#909399'))
        status_label = QLabel(status_text)
        status_label.setStyleSheet(f"color: {status_color}; border: 1px solid {status_color}; border-radius: 4px; padding: 2px 6px; font-size: 10px;")
        top_layout.addWidget(status_label)
        
        main_layout.addLayout(top_layout)
        
        # 分隔线
        line = QFrame()
        line.setFrameShape(QFrame.NoFrame)
        line.setFixedHeight(1)
        line.setProperty("class", "separator")
        main_layout.addWidget(line)
        
        # 2. 中部：合同信息
        info_layout = QGridLayout()
        info_layout.setSpacing(5)
        
        # 甲方/乙方
        info_layout.addWidget(QLabel("甲方:"), 0, 0)
        party_a = QLabel(self.data.get('party_a', '-') or '-')
        # party_a.setStyleSheet("color: #606266;") # Removed hardcoded color
        info_layout.addWidget(party_a, 0, 1)
        
        info_layout.addWidget(QLabel("乙方:"), 1, 0)
        party_b = QLabel(self.data.get('party_b', '-') or '-')
        # party_b.setStyleSheet("color: #606266;") # Removed hardcoded color
        info_layout.addWidget(party_b, 1, 1)
        
        # 金额
        info_layout.addWidget(QLabel("金额:"), 2, 0)
        amount = self.data.get('amount', 0)
        amount_val = QLabel(f"¥{amount:,.2f}")
        amount_val.setStyleSheet("font-weight: bold; color: #E6A23C;")
        info_layout.addWidget(amount_val, 2, 1)
        
        # 日期
        date_text = f"{self.data.get('signing_date', '')}"
        if self.data.get('expiration_date'):
            date_text += f" ~ {self.data.get('expiration_date', '')}"
        
        date_label = QLabel(date_text)
        date_label.setProperty("class", "info-text")
        date_label.setStyleSheet("font-size: 11px; margin-top: 5px;")
        
        main_layout.addLayout(info_layout)
        main_layout.addWidget(date_label)
        
        main_layout.addStretch()
        
        # 3. 底部：操作栏
        action_layout = QHBoxLayout()
        action_layout.setContentsMargins(0, 5, 0, 0)
        
        # 类型标签 (收款/付款)
        c_type = self.data.get('contract_type', 'incoming')
        if c_type == 'incoming' or c_type == '收款合同':
            type_text = "收款合同"
            type_icon = "💰"
        elif c_type == 'outgoing' or c_type == '付款合同':
            type_text = "付款合同"
            type_icon = "💸"
        else:
            type_text = c_type
            if "收款" in type_text:
                type_icon = "💰"
            elif "付款" in type_text:
                type_icon = "💸"
            else:
                type_icon = "📄"
        
        type_label = QLabel(f"{type_icon} {type_text}")
        type_label.setStyleSheet("color: #909399; font-size: 11px;")
        action_layout.addWidget(type_label)
        
        # 分类标签（可选）
        category_name = (self.data.get('category_name') or "").strip()
        if category_name:
            cat_label = QLabel(f"· {category_name}")
            cat_label.setStyleSheet("color: #909399; font-size: 11px;")
            action_layout.addWidget(cat_label)
        
        action_layout.addStretch()
        
        edit_btn = QPushButton("编辑")
        edit_btn.setCursor(Qt.PointingHandCursor)
        edit_btn.setProperty("class", "primary small-btn")
        edit_btn.clicked.connect(lambda: self.callbacks.get('edit', lambda: None)())
        action_layout.addWidget(edit_btn)
        
        delete_btn = QPushButton("删除")
        delete_btn.setCursor(Qt.PointingHandCursor)
        delete_btn.setProperty("class", "danger small-btn")
        delete_btn.clicked.connect(lambda: self.callbacks.get('delete', lambda: None)())
        action_layout.addWidget(delete_btn)
        
        main_layout.addLayout(action_layout)



class ContractDetailDialog(QDialog):
    """合同详情/编辑对话框"""
    def __init__(self, parent=None, db_manager=None, data=None):
        super().__init__(parent)
        self.db_manager = db_manager
        self.data = data or {}
        self.is_edit = bool(data)
        # 引用主窗口用于跨模块刷新
        try:
            self.main_window = parent.main_window if parent and hasattr(parent, 'main_window') else None
        except Exception:
            self.main_window = None
        self.setWindowTitle("编辑合同" if self.is_edit else "新增合同")
        self.setMinimumWidth(800)
        self.setMinimumHeight(600)
        self._init_ui()
    
    def _trigger_global_refresh(self):
        """刷新合同列表与首页仪表盘"""
        try:
            # 刷新合同列表
            parent = self.parent()
            if parent:
                # 刷新合同分类筛选 (新增分类后需要及时刷新)
                if hasattr(parent, '_load_category_filters'):
                    parent._load_category_filters()
                
                # 移除刷新父窗口列表的操作，防止在对话框未关闭时销毁来源CardWidget导致RuntimeError
                # if hasattr(parent, '_load_contracts'):
                #     parent._load_contracts()

            # 刷新首页仪表盘
            if self.main_window and hasattr(self.main_window, 'dashboard'):
                self.main_window.dashboard.update_data()
        except Exception as e:
            logger.error(f"Trigger global refresh failed: {e}")
        
    def _init_ui(self):
        layout = QVBoxLayout(self)
        
        # 标签页
        self.tabs = QTabWidget()
        layout.addWidget(self.tabs)
        
        # 1. 基本信息页
        self.basic_widget = QWidget()
        self._init_basic_info()
        self.tabs.addTab(self.basic_widget, "基本信息")
        
        # 只有在编辑模式下才显示附件和付款计划
        if self.is_edit:
            # 2. 附件管理页
            self.attach_widget = QWidget()
            self._init_attachments()
            self.tabs.addTab(self.attach_widget, "附件管理")
            
            # 3. 付款计划页
            self.payment_widget = QWidget()
            self._init_payment_schedule()
            self.tabs.addTab(self.payment_widget, "付款计划")
        
        # 按钮区域
        btn_layout = QHBoxLayout()
        btn_layout.addStretch()
        
        cancel_btn = QPushButton("取消")
        cancel_btn.clicked.connect(self.reject)
        btn_layout.addWidget(cancel_btn)
        
        save_btn = QPushButton("保存基本信息")
        save_btn.setProperty("class", "primary")
        save_btn.clicked.connect(self.save_data)
        btn_layout.addWidget(save_btn)
        
        layout.addLayout(btn_layout)
        
    def _check_type_usage(self, name):
        """检查合同类型使用情况"""
        count = 0
        try:
            # Direct match
            sql = "SELECT COUNT(*) FROM contracts WHERE is_deleted=0 AND contract_type=?"
            c1 = self.db_manager.execute_query(sql, (name,))[0][0]
            count += c1
            
            # Legacy mapping match
            legacy_map = {'收款合同': 'incoming', '付款合同': 'outgoing'}
            if name in legacy_map:
                legacy_val = legacy_map[name]
                c2 = self.db_manager.execute_query(sql, (legacy_val,))[0][0]
                count += c2
                
            if count > 0:
                return f"该类型已被 {count} 个合同使用，无法删除。"
        except Exception as e:
            logger.error(f"Check type usage failed: {e}")
            return f"检查使用情况失败: {e}"
        return None

    def _check_category_usage(self, name):
        """检查合同分类使用情况"""
        try:
            # Need ID for category
            res = self.db_manager.execute_query("SELECT id FROM contract_categories WHERE name=?", (name,))
            if not res: return None # Not found, safe to delete
            cat_id = res[0][0]
            
            # Check legacy category_id and multi-select category_ids
            sql = """
                SELECT COUNT(*) FROM contracts 
                WHERE is_deleted=0 AND (
                    category_id=? 
                    OR (category_ids IS NOT NULL AND ',' || category_ids || ',' LIKE ?)
                )
            """
            like_pattern = f"%,{cat_id},%"
            count = self.db_manager.execute_query(sql, (cat_id, like_pattern))[0][0]
            
            if count > 0:
                return f"该分类已被 {count} 个合同使用，无法删除。"
        except Exception as e:
            logger.error(f"Check category usage failed: {e}")
            return f"检查使用情况失败: {e}"
        return None

    def _init_basic_info(self):
        layout = QVBoxLayout(self.basic_widget)
        layout.setSpacing(20)
        layout.setContentsMargins(20, 20, 20, 20)
        
        # 1. 基本信息
        basic_group = QGroupBox("基本信息")
        basic_layout = QGridLayout()
        basic_layout.setSpacing(15)
        
        # 合同编号
        basic_layout.addWidget(QLabel("合同编号:"), 0, 0)
        self.number_input = QLineEdit(self.data.get('contract_number', ''))
        self.number_input.setPlaceholderText("系统自动生成或手动输入")
        basic_layout.addWidget(self.number_input, 0, 1)
        
        # 合同名称
        basic_layout.addWidget(QLabel("合同名称*:"), 0, 2)
        self.title_input = QLineEdit(self.data.get('title', ''))
        self.title_input.setPlaceholderText("请输入合同名称")
        basic_layout.addWidget(self.title_input, 0, 3)
        
        # 合同类型
        basic_layout.addWidget(QLabel("合同类型:"), 1, 0)
        self.type_widget = SingleSelectionWidget(
            self.db_manager, 
            'contract_types',
            check_usage_func=self._check_type_usage
        )
        current_type = self.data.get('contract_type', 'incoming')
        self.type_widget.set_selected(current_type)
        basic_layout.addWidget(self.type_widget, 1, 1)
        
        # 合同分类
        basic_layout.addWidget(QLabel("合同分类:"), 1, 2)
        self.category_widget = SingleSelectionWidget(
            self.db_manager, 
            'contract_categories',
            check_usage_func=self._check_category_usage,
            multi_select=True
        )
        # 如果是编辑模式，加载该合同的分类并选中
        if self.is_edit:
            selected_names = []
            # 1. 尝试加载 category_ids (多选)
            cat_ids_str = self.data.get('category_ids')
            if cat_ids_str:
                try:
                    id_list = [x.strip() for x in str(cat_ids_str).split(',') if x.strip()]
                    if id_list:
                        placeholders = ','.join(['?'] * len(id_list))
                        rows = self.db_manager.execute_query(f"SELECT name FROM contract_categories WHERE id IN ({placeholders})", id_list)
                        selected_names = [r[0] for r in rows]
                except Exception as e:
                    logger.error(f"Load contract category_ids failed: {e}")

            # 2. 降级兼容 category_id (单选)
            if not selected_names and self.data.get('category_id'):
                try:
                    cat_id = self.data['category_id']
                    rows = self.db_manager.execute_query("SELECT name FROM contract_categories WHERE id = ?", (cat_id,))
                    if rows and rows[0][0]:
                        selected_names.append(rows[0][0])
                except Exception as e:
                    logger.error(f"Load contract category failed: {e}")
            
            if selected_names:
                self.category_widget.set_selected(selected_names)
        basic_layout.addWidget(self.category_widget, 1, 3)
        
        basic_group.setLayout(basic_layout)
        layout.addWidget(basic_group)
        
        # 2. 签约双方
        party_group = QGroupBox("签约双方")
        party_layout = QGridLayout()
        party_layout.setSpacing(15)
        
        party_layout.addWidget(QLabel("甲方:"), 0, 0)
        self.party_a_input = QLineEdit(self.data.get('party_a', ''))
        self.party_a_input.setPlaceholderText("请输入甲方名称")
        party_layout.addWidget(self.party_a_input, 0, 1)
        
        party_layout.addWidget(QLabel("乙方:"), 0, 2)
        self.party_b_input = QLineEdit(self.data.get('party_b', ''))
        self.party_b_input.setPlaceholderText("请输入乙方名称")
        party_layout.addWidget(self.party_b_input, 0, 3)
        
        party_group.setLayout(party_layout)
        layout.addWidget(party_group)
        
        # 3. 合同详情
        detail_group = QGroupBox("合同详情")
        detail_layout = QGridLayout()
        detail_layout.setSpacing(15)
        
        # 金额
        detail_layout.addWidget(QLabel("合同金额:"), 0, 0)
        self.amount_input = QDoubleSpinBox()
        self.amount_input.setRange(0, 100000000)
        self.amount_input.setPrefix("¥ ")
        self.amount_input.setDecimals(2)
        self.amount_input.setValue(float(self.data.get('amount', 0)))
        detail_layout.addWidget(self.amount_input, 0, 1)
        
        # 状态
        detail_layout.addWidget(QLabel("合同状态:"), 0, 2)
        self.status_combo = QComboBox()
        self.status_combo.addItem("草稿", "draft")
        self.status_combo.addItem("执行中", "active")
        self.status_combo.addItem("已完成", "completed")
        self.status_combo.addItem("已过期", "expired")
        self.status_combo.addItem("已终止", "terminated")
        
        current_status = self.data.get('status', 'draft')
        index = self.status_combo.findData(current_status)
        if index >= 0:
            self.status_combo.setCurrentIndex(index)
        detail_layout.addWidget(self.status_combo, 0, 3)
        
        # 签约日期
        detail_layout.addWidget(QLabel("签约日期:"), 1, 0)
        self.sign_date_input = ModernDateEdit()
        self.sign_date_input.setDisplayFormat("yyyy-MM-dd")
        sign_date = self.data.get('signing_date')
        if sign_date:
            self.sign_date_input.setDate(QDate.fromString(sign_date, "yyyy-MM-dd"))
        else:
            self.sign_date_input.setDate(QDate.currentDate())
        detail_layout.addWidget(self.sign_date_input, 1, 1)
        
        # 到期日期
        detail_layout.addWidget(QLabel("到期日期:"), 1, 2)
        self.expire_date_input = ModernDateEdit()
        self.expire_date_input.setDisplayFormat("yyyy-MM-dd")
        expire_date = self.data.get('expiration_date')
        if expire_date:
            self.expire_date_input.setDate(QDate.fromString(expire_date, "yyyy-MM-dd"))
        else:
            self.expire_date_input.setDate(QDate.currentDate().addYears(1))
        detail_layout.addWidget(self.expire_date_input, 1, 3)
        
        detail_group.setLayout(detail_layout)
        layout.addWidget(detail_group)
        
        # 4. 备注
        remarks_group = QGroupBox("备注")
        remarks_layout = QVBoxLayout()
        self.remarks_input = QTextEdit(self.data.get('remarks', ''))
        self.remarks_input.setPlaceholderText("请输入备注信息...")
        self.remarks_input.setMaximumHeight(60)
        remarks_layout.addWidget(self.remarks_input)
        remarks_group.setLayout(remarks_layout)
        layout.addWidget(remarks_group)
        
        layout.addStretch()
    
    def _init_attachments(self):
        layout = QVBoxLayout(self.attach_widget)
        
        # 工具栏
        tool_layout = QHBoxLayout()
        
        # 生成合同按钮
        gen_btn = QPushButton("生成合同文件")
        gen_btn.setProperty("class", "success small-btn")
        gen_btn.clicked.connect(self._generate_contract_doc)
        tool_layout.addWidget(gen_btn)
        
        upload_btn = QPushButton("上传附件")
        upload_btn.setProperty("class", "primary small-btn")
        upload_btn.clicked.connect(self._upload_attachment)
        tool_layout.addWidget(upload_btn)
        
        refresh_btn = QPushButton("刷新")
        refresh_btn.setProperty("class", "small-btn")
        refresh_btn.clicked.connect(self._load_attachments)
        tool_layout.addWidget(refresh_btn)
        
        tool_layout.addStretch()
        layout.addLayout(tool_layout)
        
        # 附件列表
        self.attach_list = QTableWidget()
        self.attach_list.setColumnCount(4)
        self.attach_list.setHorizontalHeaderLabels(["文件名", "大小", "上传时间", "操作"])
        self.attach_list.horizontalHeader().setSectionResizeMode(0, QHeaderView.Stretch)
        self.attach_list.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeToContents)
        self.attach_list.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeToContents)
        self.attach_list.horizontalHeader().setSectionResizeMode(3, QHeaderView.Fixed)
        self.attach_list.setColumnWidth(3, 180)
        self.attach_list.verticalHeader().setSectionResizeMode(QHeaderView.ResizeToContents)
        self.attach_list.setSelectionMode(QAbstractItemView.NoSelection)
        self.attach_list.setEditTriggers(QAbstractItemView.NoEditTriggers)
        layout.addWidget(self.attach_list)
        
        self._load_attachments()
        
    def _load_attachments(self):
        self.attach_list.setRowCount(0)
        attachments = self.db_manager.get_contract_attachments(self.data['id'])
        
        for row in attachments:
            row_idx = self.attach_list.rowCount()
            self.attach_list.insertRow(row_idx)
            
            # 文件名
            self.attach_list.setItem(row_idx, 0, QTableWidgetItem(row[2]))
            
            # 大小
            size_kb = row[5] / 1024 if row[5] else 0
            self.attach_list.setItem(row_idx, 1, QTableWidgetItem(f"{size_kb:.1f} KB"))
            
            # 上传时间
            self.attach_list.setItem(row_idx, 2, QTableWidgetItem(row[4]))
            
            # 操作
            action_widget = QWidget()
            action_layout = QHBoxLayout(action_widget)
            action_layout.setContentsMargins(2, 2, 2, 2)
            
            open_btn = QPushButton("查看")
            open_btn.setProperty("class", "small-btn")
            open_btn.clicked.connect(lambda _, p=row[3]: self._open_attachment(p))
            action_layout.addWidget(open_btn)
            
            del_btn = QPushButton("删除")
            del_btn.setProperty("class", "danger small-btn")
            del_btn.clicked.connect(lambda _, id=row[0], p=row[3]: self._delete_attachment(id, p))
            action_layout.addWidget(del_btn)
            
            self.attach_list.setCellWidget(row_idx, 3, action_widget)

    def _generate_contract_doc(self):
        """生成合同文档"""
        if not self.data.get('id'):
            QMessageBox.warning(self, "提示", "请先保存合同基本信息")
            return

        if DocxTemplate is None:
            QMessageBox.critical(self, "错误", "缺少 docxtpl 库，无法生成合同。请安装: pip install docxtpl")
            return

        try:
            # 1. 选择模板文件（支持自定义、多模板）
            templates_dir = get_app_path('templates')
            if not os.path.exists(templates_dir):
                os.makedirs(templates_dir)
            # 如果目录里没有任何模板，自动创建一个默认模板以便用户选择
            try:
                has_docx = any(fname.lower().endswith('.docx') for fname in os.listdir(templates_dir))
            except Exception:
                has_docx = False
            if not has_docx:
                try:
                    from docx import Document
                    default_template_path = os.path.join(templates_dir, 'general_contract_template.docx')
                    if not os.path.exists(default_template_path):
                        doc = Document()
                        doc.add_heading('通用业务合同', 0)
                        doc.add_paragraph('合同编号：{{ contract_number }}')
                        doc.add_paragraph('签订日期：{{ signing_date }}')
                        doc.add_heading('甲方（委托方）：{{ party_a }}', level=1)
                        doc.add_heading('乙方（受托方）：{{ party_b }}', level=1)
                        doc.add_paragraph('甲乙双方经友好协商，就以下事项达成一致：')
                        doc.add_heading('一、合同标的', level=2)
                        doc.add_paragraph('合同名称：{{ title }}')
                        doc.add_paragraph('合同总金额：人民币 {{ amount }} 元')
                        doc.add_heading('二、有效期限', level=2)
                        doc.add_paragraph('本合同有效期自 {{ signing_date }} 至 {{ expiration_date }} 止。')
                        doc.add_heading('三、其他条款', level=2)
                        doc.add_paragraph('{{ remarks }}')
                        doc.add_paragraph('\\n')
                        doc.add_paragraph('甲方（盖章）：________________    乙方（盖章）：________________')
                        doc.add_paragraph('代表签字：__________________    代表签字：__________________')
                        doc.add_paragraph('日期：{{ signing_date }}            日期：{{ signing_date }}')
                        doc.save(default_template_path)
                        logger.info(f"Created default contract template at {default_template_path}")
                except Exception as e:
                    logger.error(f"Failed to ensure default template: {e}")
            # 弹出文件选择框，允许自由选择模板
            selected_template, _ = QFileDialog.getOpenFileName(
                self,
                "选择合同模板",
                templates_dir,
                "Word 模板 (*.docx)"
            )
            if not selected_template or not os.path.exists(selected_template):
                QMessageBox.warning(self, "错误", "未选择有效的合同模板文件")
                return
            template_path = selected_template

            # 2. 准备数据上下文
            context = {
                'contract_number': self.data.get('contract_number', ''),
                'signing_date': self.data.get('signing_date', ''),
                'expiration_date': self.data.get('expiration_date', ''),
                'party_a': self.data.get('party_a', ''),
                'party_b': self.data.get('party_b', ''),
                'title': self.data.get('title', ''),
                'amount': f"{float(self.data.get('amount', 0)):,.2f}",
                'remarks': self.data.get('remarks', '') or '无'
            }

            # 3. 生成文件
            doc = DocxTemplate(template_path)
            doc.render(context)

            # 生成目标保存路径：
            # 若甲方名称匹配客户管理中的公司名称，则保存到客户资料路径/公司名 目录下
            # 否则保存到 data/contract_attachments
            party_a_name = (self.data.get('party_a') or '').strip()
            target_dir = None
            try:
                from PyQt5.QtCore import QSettings
                settings = QSettings("CustomerManagement", "CustomerWindow")
                customer_root = (settings.value("customer_data_path", "") or "").strip()
                customer_exists = False
                if customer_root and party_a_name:
                    try:
                        res = self.db_manager.execute_query(
                            "SELECT 1 FROM customers WHERE is_deleted = 0 AND company_name = ?",
                            (party_a_name,)
                        )
                        customer_exists = bool(res)
                    except Exception as e:
                        logger.error(f"Check customer existence failed: {e}")
                if customer_root and customer_exists:
                    target_dir = os.path.join(customer_root, party_a_name)
                else:
                    target_dir = os.path.join(get_app_path('data'), 'contract_attachments')
            except Exception as e:
                logger.error(f"Resolve target dir failed: {e}")
                target_dir = os.path.join(get_app_path('data'), 'contract_attachments')
            if not os.path.exists(target_dir):
                os.makedirs(target_dir)

            # 文件名：XX公司-合同；若无甲方名称则退回原命名规则
            base_name = f"{party_a_name}-合同" if party_a_name else f"合同_{self.data.get('contract_number', 'new')}"
            file_name = f"{base_name}.docx"
            target_path = os.path.join(target_dir, file_name)
            if os.path.exists(target_path):
                timestamp = datetime.now().strftime('%Y%m%d%H%M')
                file_name = f"{base_name}-{timestamp}.docx"
                target_path = os.path.join(target_dir, file_name)
            doc.save(target_path)

            # 4. 添加到附件记录
            file_size = os.path.getsize(target_path)
            attach_data = {
                'contract_id': self.data['id'],
                'file_name': file_name,
                'file_path': target_path,
                'file_size': file_size
            }
            
            # 调用 db_manager.add_contract_attachment
            if self.db_manager.add_contract_attachment(attach_data):
                self._load_attachments()
                self._trigger_global_refresh()
                QMessageBox.information(self, "成功", f"合同文件已生成并添加至附件:\n{file_name}")
                
                # 询问是否打开
                reply = QMessageBox.question(self, "提示", "是否立即打开生成的合同文件？", QMessageBox.Yes | QMessageBox.No)
                if reply == QMessageBox.Yes:
                    self._open_attachment(target_path)
            else:
                QMessageBox.warning(self, "错误", "文件已生成但保存记录失败")

        except Exception as e:
            logger.error(f"Generate contract failed: {e}")
            QMessageBox.critical(self, "错误", f"生成失败: {str(e)}")

    def _upload_attachment(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "选择文件")
        if not file_path:
            return
            
        try:
            # 准备目标路径
            file_name = os.path.basename(file_path)
            # 使用 contract_attachments 目录
            target_dir = os.path.join(get_app_path('data'), 'contract_attachments')
            if not os.path.exists(target_dir):
                os.makedirs(target_dir)
                
            # 生成唯一文件名防止覆盖
            timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
            target_name = f"{timestamp}_{file_name}"
            target_path = os.path.join(target_dir, target_name)
            
            # 复制文件
            shutil.copy2(file_path, target_path)
            
            # 记录到数据库
            file_size = os.path.getsize(target_path)
            data = {
                'contract_id': self.data['id'],
                'file_name': file_name,
                'file_path': target_path,
                'file_size': file_size
            }
            
            if self.db_manager.add_contract_attachment(data):
                self._load_attachments()
                self._trigger_global_refresh()
            else:
                QMessageBox.warning(self, "错误", "保存附件记录失败")
                
        except Exception as e:
            logger.error(f"Upload attachment failed: {e}")
            QMessageBox.critical(self, "错误", f"上传失败: {str(e)}")

    def _open_attachment(self, path):
        try:
            if os.path.exists(path):
                QDesktopServices.openUrl(QUrl.fromLocalFile(path))
            else:
                QMessageBox.warning(self, "错误", "文件不存在，可能已被删除")
        except Exception as e:
            QMessageBox.critical(self, "错误", f"打开文件失败: {str(e)}")

    def _delete_attachment(self, attach_id, path):
        reply = QMessageBox.question(self, "确认", "确定要删除此附件吗？", QMessageBox.Yes | QMessageBox.No)
        if reply != QMessageBox.Yes:
            return
            
        if self.db_manager.delete_contract_attachment(attach_id):
            # 尝试删除物理文件
            try:
                if os.path.exists(path):
                    os.remove(path)
            except Exception as e:
                logger.error(f"Failed to delete file {path}: {e}")
            
            self._load_attachments()
            self._trigger_global_refresh()
        else:
            QMessageBox.warning(self, "错误", "删除失败")

    def _init_payment_schedule(self):
        layout = QVBoxLayout(self.payment_widget)
        
        # 工具栏
        tool_layout = QHBoxLayout()
        add_btn = QPushButton("添加分期")
        add_btn.setProperty("class", "primary small-btn")
        add_btn.clicked.connect(self._add_payment_schedule)
        tool_layout.addWidget(add_btn)
        
        refresh_btn = QPushButton("刷新")
        refresh_btn.setProperty("class", "small-btn")
        refresh_btn.clicked.connect(self._load_payment_schedules)
        tool_layout.addWidget(refresh_btn)
        
        tool_layout.addStretch()
        layout.addLayout(tool_layout)
        
        # 列表
        self.payment_list = QTableWidget()
        self.payment_list.setColumnCount(5)
        self.payment_list.setHorizontalHeaderLabels(["期数", "应付/应收日期", "金额", "状态", "操作"])
        self.payment_list.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeToContents)
        self.payment_list.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeToContents)
        self.payment_list.horizontalHeader().setSectionResizeMode(2, QHeaderView.Stretch)
        self.payment_list.horizontalHeader().setSectionResizeMode(3, QHeaderView.ResizeToContents)
        self.payment_list.horizontalHeader().setSectionResizeMode(4, QHeaderView.Fixed)
        self.payment_list.setColumnWidth(4, 200)
        self.payment_list.verticalHeader().setSectionResizeMode(QHeaderView.ResizeToContents)
        self.payment_list.setSelectionMode(QAbstractItemView.NoSelection)
        self.payment_list.setEditTriggers(QAbstractItemView.NoEditTriggers)
        layout.addWidget(self.payment_list)
        
        self._load_payment_schedules()

    def _load_payment_schedules(self):
        self.payment_list.setRowCount(0)
        schedules = self.db_manager.get_payment_schedules(self.data['id'])
        
        for row in schedules:
            row_idx = self.payment_list.rowCount()
            self.payment_list.insertRow(row_idx)
            
            # 期数
            self.payment_list.setItem(row_idx, 0, QTableWidgetItem(f"第 {row[2]} 期"))
            
            # 日期
            self.payment_list.setItem(row_idx, 1, QTableWidgetItem(row[3]))
            
            # 金额
            amount_item = QTableWidgetItem(f"¥{row[4]:,.2f}")
            amount_item.setTextAlignment(Qt.AlignRight | Qt.AlignVCenter)
            self.payment_list.setItem(row_idx, 2, amount_item)
            
            # 状态
            status_map = {'pending': '未结清', 'paid': '已结清'}
            status_text = status_map.get(row[5], row[5])
            self.payment_list.setItem(row_idx, 3, QTableWidgetItem(status_text))
            
            # 操作
            action_widget = QWidget()
            action_layout = QHBoxLayout(action_widget)
            action_layout.setContentsMargins(2, 2, 2, 2)
            
            toggle_btn = QPushButton("标记结清" if row[5] == 'pending' else "标记未付")
            toggle_btn.setProperty("class", "small-btn")
            toggle_btn.setMinimumWidth(90)  # 防止文字截断
            toggle_btn.clicked.connect(lambda _, id=row[0], s=row[5]: self._toggle_payment_status(id, s))
            action_layout.addWidget(toggle_btn)
            
            del_btn = QPushButton("删除")
            del_btn.setProperty("class", "danger small-btn")
            del_btn.clicked.connect(lambda _, id=row[0]: self._delete_payment_schedule(id))
            action_layout.addWidget(del_btn)
            
            self.payment_list.setCellWidget(row_idx, 4, action_widget)

    def _add_payment_schedule(self):
        dialog = QDialog(self)
        dialog.setWindowTitle("添加分期计划")
        dialog.setMinimumWidth(300)
        layout = QFormLayout(dialog)
        
        num_input = QDoubleSpinBox()
        num_input.setDecimals(0)
        num_input.setRange(1, 100)
        # 自动推断下一期
        current_rows = self.payment_list.rowCount()
        num_input.setValue(current_rows + 1)
        layout.addRow("期数:", num_input)
        
        date_input = ModernDateEdit()
        date_input.setDate(QDate.currentDate().addMonths(1))
        date_input.setDisplayFormat("yyyy-MM-dd")
        layout.addRow("日期:", date_input)
        
        amount_input = QDoubleSpinBox()
        amount_input.setRange(0, 100000000)
        amount_input.setPrefix("¥ ")
        # 尝试推断剩余金额
        total_amount = float(self.data.get('amount', 0))
        # 这里简单处理
        layout.addRow("金额:", amount_input)
        
        btn_box = QHBoxLayout()
        ok_btn = QPushButton("确定")
        ok_btn.clicked.connect(dialog.accept)
        cancel_btn = QPushButton("取消")
        cancel_btn.clicked.connect(dialog.reject)
        btn_box.addWidget(ok_btn)
        btn_box.addWidget(cancel_btn)
        layout.addRow(btn_box)
        
        if dialog.exec_() == QDialog.Accepted:
            data = {
                'contract_id': self.data['id'],
                'installment_number': int(num_input.value()),
                'due_date': date_input.date().toString("yyyy-MM-dd"),
                'amount': amount_input.value(),
                'status': 'pending'
            }
            if self.db_manager.add_payment_schedule(data):
                self._load_payment_schedules()
                self._trigger_global_refresh()

    def _toggle_payment_status(self, schedule_id, current_status):
        new_status = 'paid' if current_status == 'pending' else 'pending'
        if self.db_manager.update_payment_schedule(schedule_id, {'status': new_status}):
            self._load_payment_schedules()
            self._trigger_global_refresh()

    def _delete_payment_schedule(self, schedule_id):
        reply = QMessageBox.question(self, "确认", "确定要删除此分期计划吗？", QMessageBox.Yes | QMessageBox.No)
        if reply == QMessageBox.Yes:
            if self.db_manager.delete_payment_schedule(schedule_id):
                self._load_payment_schedules()
                self._trigger_global_refresh()

    def save_data(self):
        """保存合同基本信息"""
        try:
            # 1. 验证必填项
            title = self.title_input.text().strip()
            if not title:
                QMessageBox.warning(self, "提示", "合同名称不能为空")
                return

            # 2. 获取数据
            contract_type = self.type_widget.get_selected()
            if not contract_type:
                QMessageBox.warning(self, "提示", "请选择合同类型")
                return
            
            # Map back to legacy if needed, or keep as text
            if contract_type == '收款合同': contract_type = 'incoming'
            elif contract_type == '付款合同': contract_type = 'outgoing'
            
            # 分类处理 (多选)
            selected_names = self.category_widget.get_selected()
            if isinstance(selected_names, str):
                selected_names = [selected_names] if selected_names else []
            
            category_ids_list = []
            for name in selected_names:
                try:
                    res = self.db_manager.execute_query("SELECT id FROM contract_categories WHERE name=?", (name,))
                    if res:
                        cid = res[0][0]
                    else:
                        self.db_manager.execute_query("INSERT INTO contract_categories (name) VALUES (?)", (name,), fetch=False)
                        self.db_manager.conn.commit()
                        res = self.db_manager.execute_query("SELECT last_insert_rowid()")
                        cid = res[0][0]
                    category_ids_list.append(str(cid))
                except Exception as e:
                    logger.error(f"Handle category {name} failed: {e}")

            category_ids_str = ",".join(category_ids_list)
            category_id = int(category_ids_list[0]) if category_ids_list else None
            
            contract_data = {
                'contract_number': self.number_input.text().strip() or f"HT{datetime.now().strftime('%Y%m%d%H%M%S')}",
                'title': title,
                'contract_type': contract_type,
                'category_id': category_id,
                'category_ids': category_ids_str,
                'party_a': self.party_a_input.text().strip(),
                'party_b': self.party_b_input.text().strip(),
                'amount': self.amount_input.value(),
                'signing_date': self.sign_date_input.date().toString("yyyy-MM-dd"),
                'expiration_date': self.expire_date_input.date().toString("yyyy-MM-dd"),
                'status': self.status_combo.currentData(),
                'remarks': self.remarks_input.toPlainText().strip()
            }
            
            if self.is_edit:
                # 更新
                sql = """
                    UPDATE contracts SET 
                        contract_number=?, title=?, contract_type=?, 
                        category_id=?, category_ids=?, party_a=?, party_b=?, amount=?, 
                        signing_date=?, expiration_date=?, status=?, remarks=?,
                        updated_at=CURRENT_TIMESTAMP
                    WHERE id=?
                """
                params = (
                    contract_data['contract_number'], contract_data['title'], contract_data['contract_type'],
                    contract_data.get('category_id'),
                    contract_data.get('category_ids'),
                    contract_data['party_a'], contract_data['party_b'], contract_data['amount'],
                    contract_data['signing_date'], contract_data['expiration_date'], contract_data['status'],
                    contract_data['remarks'], self.data['id']
                )
            else:
                # 新增
                sql = """
                    INSERT INTO contracts (
                        contract_number, title, contract_type, 
                        category_id, category_ids, party_a, party_b, amount, 
                        signing_date, expiration_date, status, remarks
                    ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """
                params = (
                    contract_data['contract_number'], contract_data['title'], contract_data['contract_type'],
                    contract_data.get('category_id'),
                    contract_data.get('category_ids'),
                    contract_data['party_a'], contract_data['party_b'], contract_data['amount'],
                    contract_data['signing_date'], contract_data['expiration_date'], contract_data['status'],
                    contract_data['remarks']
                )
                
            self.db_manager.execute_query(sql, params, fetch=False)
            self.db_manager.conn.commit()
            
            if not self.is_edit:
                reply = QMessageBox.question(self, "成功", "合同已保存，是否继续添加附件或付款计划？", 
                                           QMessageBox.Yes | QMessageBox.No)
                if reply == QMessageBox.Yes:
                    cursor = self.db_manager.execute_query("SELECT last_insert_rowid()")
                    new_id = cursor[0][0]
                    self.data = contract_data
                    self.data['id'] = new_id
                    self.is_edit = True
                    self.setWindowTitle("编辑合同")
                    
                    # 移除旧的布局并重新初始化（简单粗暴）
                    QWidget().setLayout(self.layout()) # 清除引用
                    self._init_ui()
                    self.tabs.setCurrentIndex(1) # 跳转到附件页
                else:
                    self.accept()
            else:
                QMessageBox.information(self, "成功", "保存成功")
            
            # 全局刷新
            self._trigger_global_refresh()
            
        except Exception as e:
            logger.error(f"Save contract failed: {e}")
            QMessageBox.critical(self, "错误", f"保存失败: {str(e)}")

class ContractWindow(QWidget):
    def __init__(self, db_manager, parent=None):
        super().__init__(parent)
        self.db_manager = db_manager
        # 保存主窗口引用用于跨模块刷新
        self.main_window = parent
        
        # Pagination state
        self.page = 1
        self.page_size = 20
        self.total_pages = 1
        self.total_count = 0
        self.pending_select_query = None
        self.threadpool = QThreadPool()
        
        self._init_ui()
        self._load_contracts()
        
    def _init_ui(self):
        main_layout = QVBoxLayout()
        main_layout.setContentsMargins(20, 20, 20, 20)
        main_layout.setSpacing(20)
        
        # 顶部操作栏
        top_frame = QFrame()
        top_frame.setProperty("class", "card")
        top_layout = QVBoxLayout(top_frame)
        top_layout.setSpacing(10)
        
        # 搜索行
        row1 = QHBoxLayout()
        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText('搜索合同名称、编号或当事人...')
        self.search_input.textChanged.connect(self._on_search_changed)
        row1.addWidget(self.search_input)
        top_layout.addLayout(row1)
        
        # 筛选和按钮行
        row2 = QHBoxLayout()
        row2.setSpacing(15)
        
        self.type_filter = QComboBox()
        self.type_filter.addItems(["所有类型", "收款合同", "付款合同"])
        self.type_filter.currentIndexChanged.connect(self._on_filter_changed)
        row2.addWidget(QLabel("合同类型:"))
        row2.addWidget(self.type_filter)
        
        # 合同分类筛选
        self.category_filter = QComboBox()
        self._category_id_by_name = {}
        self._load_category_filters()
        self.category_filter.currentIndexChanged.connect(self._on_filter_changed)
        row2.addWidget(QLabel("合同分类:"))
        row2.addWidget(self.category_filter)
        
        self.status_filter = QComboBox()
        self.status_filter.addItems(["所有状态", "草稿", "执行中", "已完成", "已过期", "已终止"])
        self.status_filter.currentIndexChanged.connect(self._on_filter_changed)
        row2.addWidget(QLabel("合同状态:"))
        row2.addWidget(self.status_filter)
        
        # 排序下拉框
        self.sort_combo = QComboBox()
        self.sort_combo.addItems([
            '默认排序 (创建时间)',
            '创建时间 (新→旧)',
            '创建时间 (旧→新)',
            '合同名称 (A→Z)',
            '合同名称 (Z→A)',
            '金额 (高→低)',
            '金额 (低→高)'
        ])
        # 设置下拉列表的最小宽度
        self.sort_combo.view().setMinimumWidth(160)
        self.sort_combo.currentIndexChanged.connect(self._on_sort_changed)
        row2.addWidget(QLabel("排序:"))
        row2.addWidget(self.sort_combo)

        row2.addStretch()
        
        add_btn = QPushButton("新增合同")
        add_btn.setProperty("class", "success")
        add_btn.clicked.connect(self._show_add_dialog)
        row2.addWidget(add_btn)
        
        refresh_btn = QPushButton("刷新")
        refresh_btn.setProperty("class", "primary")
        refresh_btn.clicked.connect(self._load_contracts)
        row2.addWidget(refresh_btn)
        
        top_layout.addLayout(row2)
        main_layout.addWidget(top_frame)
        
        # 列表区域
        list_frame = QFrame()
        list_frame.setProperty("class", "card")
        list_layout = QVBoxLayout(list_frame)
        list_layout.setContentsMargins(12, 12, 12, 12)
        
        batch_layout = QHBoxLayout()
        select_all_btn = QPushButton("全选")
        select_all_btn.clicked.connect(self._select_all_contracts)
        invert_btn = QPushButton("反选")
        invert_btn.clicked.connect(self._invert_selection_contracts)
        clear_btn = QPushButton("取消选择")
        clear_btn.clicked.connect(self._clear_selection_contracts)
        delete_selected_btn = QPushButton("删除选中")
        delete_selected_btn.setProperty("class", "danger")
        delete_selected_btn.clicked.connect(self._delete_selected_contracts)
        batch_layout.addWidget(select_all_btn)
        batch_layout.addWidget(invert_btn)
        batch_layout.addWidget(clear_btn)
        batch_layout.addWidget(delete_selected_btn)
        list_layout.addLayout(batch_layout)
        
        self.list_widget = QListWidget()
        self.list_widget.setFrameShape(QFrame.NoFrame)
        self.list_widget.setVerticalScrollMode(QAbstractItemView.ScrollPerPixel)
        self.list_widget.setSpacing(12)
        self.list_widget.setViewMode(QListWidget.IconMode)
        self.list_widget.setResizeMode(QListWidget.Adjust)
        self.list_widget.setMovement(QListWidget.Static)
        self.list_widget.setSelectionMode(QAbstractItemView.NoSelection)
        self.list_widget.setGridSize(QSize(372, 266)) # 略大于卡片尺寸
        self.list_widget.setStyleSheet("QListWidget { background: transparent; }")
        
        list_layout.addWidget(self.list_widget)
        
        # Pagination Controls
        pagination_layout = QHBoxLayout()
        self.prev_btn = QPushButton("上一页")
        self.prev_btn.setFixedWidth(80)
        self.prev_btn.clicked.connect(self._prev_page)
        
        self.next_btn = QPushButton("下一页")
        self.next_btn.setFixedWidth(80)
        self.next_btn.clicked.connect(self._next_page)
        
        self.page_label = QLabel("第 1 页 / 共 1 页")
        self.page_label.setAlignment(Qt.AlignCenter)
        
        # 跳转控件
        jump_container = QWidget()
        jump_container.setSizePolicy(QSizePolicy.Fixed, QSizePolicy.Preferred)
        jump_layout = QHBoxLayout(jump_container)
        jump_layout.setContentsMargins(0, 0, 0, 0)
        jump_layout.setSpacing(2)
        
        lbl_jump = QLabel("跳转至")
        lbl_jump.setSizePolicy(QSizePolicy.Fixed, QSizePolicy.Preferred)
        jump_layout.addWidget(lbl_jump)
        
        self.jump_input = QLineEdit()
        self.jump_input.setFixedSize(50, 26)
        self.jump_input.setAlignment(Qt.AlignCenter)
        self.jump_input.setValidator(QIntValidator(1, 9999))
        self.jump_input.returnPressed.connect(self._jump_to_page)
        jump_layout.addWidget(self.jump_input)
        
        lbl_page = QLabel("页")
        lbl_page.setSizePolicy(QSizePolicy.Fixed, QSizePolicy.Preferred)
        jump_layout.addWidget(lbl_page)
        
        jump_btn = QPushButton("Go")
        jump_btn.setFixedSize(30, 26)
        jump_btn.setCursor(Qt.PointingHandCursor)
        jump_btn.setStyleSheet("QPushButton { min-width: 30px; max-width: 30px; padding: 0px; margin: 0px; }")
        jump_btn.clicked.connect(self._jump_to_page)
        jump_layout.addWidget(jump_btn)

        pagination_layout.addStretch()
        pagination_layout.addWidget(self.prev_btn)
        pagination_layout.addWidget(self.page_label)
        pagination_layout.addWidget(self.next_btn)
        pagination_layout.addSpacing(5)
        pagination_layout.addWidget(jump_container)
        pagination_layout.addStretch()
        
        list_layout.addLayout(pagination_layout)
        
        main_layout.addWidget(list_frame)
        
        self.setLayout(main_layout)
        
    def _fetch_data_worker(self, search_text, type_filter, category_filter_id, status_filter, sort_option, limit, offset):
        """Background worker to fetch data and count"""
        conn = self.db_manager.create_new_connection()
        try:
            cursor = conn.cursor()
            
            where_clauses = ["c.is_deleted = 0"]
            params = []
            
            if search_text:
                where_clauses.append("(c.title LIKE ? OR c.contract_number LIKE ? OR c.party_a LIKE ? OR c.party_b LIKE ?)")
                p = f"%{search_text}%"
                params.extend([p, p, p, p])
            
            if type_filter != "所有类型":
                t = 'incoming' if type_filter == "收款合同" else 'outgoing'
                where_clauses.append("c.contract_type = ?")
                params.append(t)
            
            if category_filter_id is not None:
                # Support both legacy category_id and new category_ids
                # For category_ids (text), we check if it contains the ID
                # We wrap with commas to ensure exact match (e.g. ",1," matches ",1,2," but not ",11,")
                # SQLite concatenation is ||
                where_clauses.append("(c.category_id = ? OR (',' || IFNULL(c.category_ids, '') || ',') LIKE ?)")
                params.append(category_filter_id)
                params.append(f"%,{category_filter_id},%")
            
            if status_filter != "所有状态":
                status_map = {
                    "草稿": "draft",
                    "执行中": "active",
                    "已完成": "completed",
                    "已过期": "expired",
                    "已终止": "terminated"
                }
                s = status_map.get(status_filter)
                if s:
                    where_clauses.append("c.status = ?")
                    params.append(s)
            
            where_sql = " AND ".join(where_clauses)
            
            # Count
            count_sql = f"SELECT COUNT(*) FROM contracts c WHERE {where_sql}"
            cursor.execute(count_sql, params)
            total = cursor.fetchone()[0]
            
            # Fetch Category Map (small enough to fetch every time or could be cached if static, but safer here)
            cat_map = {}
            cursor.execute("SELECT id, name FROM contract_categories")
            for r in cursor.fetchall():
                cat_map[str(r[0])] = r[1]

            # Determine Order
            order_by = "c.created_at DESC"
            if sort_option == '创建时间 (新→旧)':
                order_by = "c.created_at DESC"
            elif sort_option == '创建时间 (旧→新)':
                order_by = "c.created_at ASC"
            elif sort_option == '合同名称 (A→Z)':
                order_by = "c.title ASC"
            elif sort_option == '合同名称 (Z→A)':
                order_by = "c.title DESC"
            elif sort_option == '金额 (高→低)':
                order_by = "c.amount DESC"
            elif sort_option == '金额 (低→高)':
                order_by = "c.amount ASC"
                
            # Data
            # Note: We don't join contract_categories here because we handle it via cat_map for both legacy and new fields
            sql = f"""
                SELECT c.id, c.contract_number, c.title, c.contract_type, 
                       c.party_a, c.party_b, c.signing_date, c.expiration_date, 
                       c.amount, c.status, c.remarks, c.category_id,
                       c.category_ids
                FROM contracts c
                WHERE {where_sql}
                ORDER BY {order_by}
                LIMIT ? OFFSET ?
            """
            cursor.execute(sql, params + [limit, offset])
            rows = cursor.fetchall()
            
            return rows, total, cat_map
        finally:
            conn.close()

    def _load_contracts(self):
        """Async load contracts with pagination"""
        # UI state
        self.prev_btn.setEnabled(False)
        self.next_btn.setEnabled(False)
        self.list_widget.clear()
        self._cards = []
        
        # Parameters
        search_text = self.search_input.text().strip()
        type_filter = self.type_filter.currentText()
        category_filter = self.category_filter.currentText()
        category_filter_id = self._category_id_by_name.get(category_filter) if category_filter != "所有分类" else None
        status_filter = self.status_filter.currentText()
        sort_option = self.sort_combo.currentText()
        offset = (self.page - 1) * self.page_size
        
        # Start worker
        worker = Worker(self._fetch_data_worker, search_text, type_filter, category_filter_id, status_filter, sort_option, self.page_size, offset)
        worker.signals.result.connect(self._on_load_success)
        worker.signals.error.connect(self._on_load_error)
        self.threadpool.start(worker)

    def _on_sort_changed(self):
        self.page = 1
        self._load_contracts()

    def _jump_to_page(self):
        text = self.jump_input.text().strip()
        if not text:
            return
        try:
            page = int(text)
            if 1 <= page <= self.total_pages:
                self.page = page
                self._load_contracts()
                self.jump_input.clear()
            else:
                QMessageBox.warning(self, "提示", f"请输入 1 到 {self.total_pages} 之间的页码")
                self.jump_input.selectAll()
                self.jump_input.setFocus()
        except ValueError:
            pass

    def _on_load_success(self, result):
        rows, total, cat_map = result
        self.total_count = total
        self.total_pages = max(1, (total + self.page_size - 1) // self.page_size)
        
        # Ensure page is within bounds
        if self.page > self.total_pages:
            self.page = self.total_pages
            if self.total_count > 0:
                self._load_contracts()
                return

        for row in rows:
            # Resolve category names
            cat_ids_str = row[12]
            legacy_cat_id = row[11]
            display_cat_name = ""
            
            # Prioritize multi-select field, fall back to legacy if needed (or combine?)
            # Logic in original code:
            # if cat_ids_str: parse and join names
            # else: display_cat_name = row[12] (from join)
            
            found_names = []
            if cat_ids_str:
                try:
                    ids = [x.strip() for x in str(cat_ids_str).split(',') if x.strip()]
                    found_names = [cat_map.get(cid, "") for cid in ids if cid in cat_map]
                except Exception:
                    pass
            
            if not found_names and legacy_cat_id:
                 name = cat_map.get(str(legacy_cat_id))
                 if name:
                     found_names.append(name)
            
            if found_names:
                display_cat_name = ", ".join(found_names)

            data = {
                'id': row[0],
                'contract_number': row[1],
                'title': row[2],
                'contract_type': row[3],
                'party_a': row[4],
                'party_b': row[5],
                'signing_date': row[6],
                'expiration_date': row[7],
                'amount': row[8],
                'status': row[9],
                'remarks': row[10],
                'category_id': legacy_cat_id,
                'category_name': display_cat_name,
                'category_ids': cat_ids_str
            }
            
            card = ContractCardWidget(data)
            card.set_callback('edit', lambda d=data: self._show_edit_dialog(d))
            card.set_callback('delete', lambda d=data: self._delete_contract(d))
            
            item = QListWidgetItem(self.list_widget)
            item.setSizeHint(QSize(360, 250))
            self.list_widget.setItemWidget(item, card)
            self._cards.append(card)
            
        self._update_pagination_ui()
        
        # 处理待处理的选中请求
        if self.pending_select_query:
            query = self.pending_select_query
            self.pending_select_query = None
            
            # 尝试查找匹配项
            for i in range(self.list_widget.count()):
                item = self.list_widget.item(i)
                card = self.list_widget.itemWidget(item)
                if card:
                    data = card.data
                    # 匹配标题、合同号、甲方、乙方
                    match_fields = [
                        data.get('title', ''),
                        data.get('contract_number', ''),
                        data.get('party_a', ''),
                        data.get('party_b', '')
                    ]
                    if any(query in str(f) for f in match_fields):
                        item.setSelected(True)
                        self.list_widget.setCurrentItem(item)
                        self.list_widget.scrollToItem(item)
                        # 触发选中视觉效果(如果需要)
                        if hasattr(card, 'set_checked'):
                            card.set_checked(True)
                        break

    def _on_load_error(self, err):
        logger.error(f"Load contracts failed: {err}")
        QMessageBox.critical(self, "Error", f"Failed to load contracts: {err}")
        self._update_pagination_ui()

    def _update_pagination_ui(self):
        self.page_label.setText(f"第 {self.page} 页 / 共 {self.total_pages} 页")
        self.prev_btn.setEnabled(self.page > 1)
        self.next_btn.setEnabled(self.page < self.total_pages)

    def _prev_page(self):
        if self.page > 1:
            self.page -= 1
            self._load_contracts()

    def _next_page(self):
        if self.page < self.total_pages:
            self.page += 1
            self._load_contracts()

    def search_and_select(self, query):
        """外部调用搜索并选中第一条"""
        self.pending_select_query = query
        # 重置筛选
        self.type_filter.setCurrentIndex(0)
        self.category_filter.setCurrentIndex(0)
        self.status_filter.setCurrentIndex(0)
        
        self.search_input.setText(query)
        # textChanged triggers _on_search_changed -> _load_contracts

    def _on_search_changed(self):
        self.page = 1
        self._load_contracts()

    def _on_filter_changed(self):
        self.page = 1
        self._load_contracts()
        
    def _apply_filters(self):
        # Legacy method kept for compatibility if called from outside, redirects to reload
        self.page = 1
        self._load_contracts()
    
    def _load_category_filters(self):
        """加载分类筛选项"""
        try:
            current_text = self.category_filter.currentText()
            # 阻塞信号，防止清除时触发 currentIndexChanged 导致 _load_contracts 重建列表
            self.category_filter.blockSignals(True)
            
            self.category_filter.clear()
            self.category_filter.addItem("所有分类")
            self._category_id_by_name = {}
            rows = []
            if hasattr(self.db_manager, 'get_contract_categories'):
                rows = self.db_manager.get_contract_categories()
            else:
                rows = self.db_manager.execute_query("SELECT id, name FROM contract_categories ORDER BY name")
            for r in rows or []:
                name = r[1]
                cid = int(r[0])
                if name:
                    self.category_filter.addItem(name)
                    self._category_id_by_name[name] = cid
            
            # 恢复选中项
            idx = self.category_filter.findText(current_text)
            if idx >= 0:
                self.category_filter.setCurrentIndex(idx)
            else:
                self.category_filter.setCurrentIndex(0)
                
            self.category_filter.blockSignals(False)
            
        except Exception as e:
            logger.error(f"Load category filters failed: {e}")
            self.category_filter.blockSignals(False)
            
    def _show_add_dialog(self):
        dialog = ContractDetailDialog(self, self.db_manager)
        if dialog.exec_() == QDialog.Accepted:
            self._load_contracts()
            
    def _show_edit_dialog(self, data):
        dialog = ContractDetailDialog(self, self.db_manager, data)
        if dialog.exec_() == QDialog.Accepted:
            self._load_contracts()
            
    def _delete_contract(self, data):
        reply = QMessageBox.question(
            self, "确认删除", 
            f"确定要删除合同 '{data.get('title')}' 吗？\n删除后可在回收站恢复。",
            QMessageBox.Yes | QMessageBox.No
        )
        
        if reply == QMessageBox.Yes:
            if self.db_manager.soft_delete_record('contracts', data['id']):
                self._load_contracts()
                # 同步刷新首页
                if hasattr(self, 'main_window') and self.main_window and hasattr(self.main_window, 'dashboard'):
                    try:
                        self.main_window.dashboard.update_data()
                    except Exception as e:
                        logger.error(f"Dashboard refresh failed: {e}")
            else:
                QMessageBox.warning(self, "错误", "删除失败")
    
    def _visible_cards(self):
        result = []
        for i in range(self.list_widget.count()):
            item = self.list_widget.item(i)
            if item.isHidden():
                continue
            card = self.list_widget.itemWidget(item)
            if card:
                result.append(card)
        return result
    
    def _select_all_contracts(self):
        for card in self._visible_cards():
            card.set_checked(True)
    
    def _invert_selection_contracts(self):
        for card in self._visible_cards():
            card.set_checked(not card.is_checked())
    
    def _clear_selection_contracts(self):
        for card in self._visible_cards():
            card.set_checked(False)
    
    def _delete_selected_contracts(self):
        targets = []
        for i in range(self.list_widget.count()):
            item = self.list_widget.item(i)
            if item.isHidden():
                continue
            card = self.list_widget.itemWidget(item)
            if card and card.is_checked():
                targets.append(card.data)
        if not targets:
            QMessageBox.information(self, "提示", "未选择任何合同")
            return
        reply = QMessageBox.question(
            self,
            "确认删除",
            f"确定要删除选中的 {len(targets)} 个合同吗？\n删除后可在回收站恢复。",
            QMessageBox.Yes | QMessageBox.No
        )
        if reply != QMessageBox.Yes:
            return
        ok = True
        for d in targets:
            if not self.db_manager.soft_delete_record('contracts', d['id']):
                ok = False
        self._load_contracts()
        if hasattr(self, 'main_window') and self.main_window and hasattr(self.main_window, 'dashboard'):
            try:
                self.main_window.dashboard.update_data()
            except Exception as e:
                logger.error(f"Dashboard refresh failed: {e}")
        if not ok:
            QMessageBox.warning(self, "错误", "部分合同删除失败")
